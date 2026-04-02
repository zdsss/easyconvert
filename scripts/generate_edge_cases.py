#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成边缘情况测试简历
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class EdgeCaseGenerator:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(self):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        return doc

    def add_heading(self, doc, text, level=1):
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading

    def save_document(self, doc, filename):
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        print(f"✓ 已生成: {filename}")

    def generate_edge_11_missing_info(self):
        """缺失信息简历"""
        doc = self.create_document()

        title = doc.add_heading('赵敏', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('软件工程师').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('某科技公司 | 工程师 | 2020-2023')
        doc.add_paragraph('• 负责后端开发')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('某大学 | 计算机科学 | 2016-2020')

        self.save_document(doc, 'edge-11-missing-info.docx')

    def generate_edge_12_date_formats(self):
        """日期格式变体"""
        doc = self.create_document()

        title = doc.add_heading('孙悟空', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('电话: 13900139999 | 邮箱: sunwukong@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('公司A | 职位A | 2020/01 - 2022/12')
        doc.add_paragraph('• 工作内容描述')
        doc.add_paragraph()

        doc.add_paragraph('公司B | 职位B | 2018.3 - 至今')
        doc.add_paragraph('• 工作内容描述')
        doc.add_paragraph()

        doc.add_paragraph('Company C | Position C | Jan 2016 - Present')
        doc.add_paragraph('• Job description')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('某大学 | 本科 | 2012-2016')

        self.save_document(doc, 'edge-12-date-formats.docx')

    def generate_edge_13_multiple_positions(self):
        """多职位同公司"""
        doc = self.create_document()

        title = doc.add_heading('周杰伦', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('电话: 13800138888 | 邮箱: jay@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('腾讯科技 | 技术总监 | 2022.01-至今')
        doc.add_paragraph('• 管理技术团队')
        doc.add_paragraph()

        doc.add_paragraph('腾讯科技 | 高级工程师 | 2019.06-2021.12')
        doc.add_paragraph('• 负责核心系统开发')
        doc.add_paragraph()

        doc.add_paragraph('腾讯科技 | 工程师 | 2017.07-2019.05')
        doc.add_paragraph('• 参与项目开发')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('台湾大学 | 音乐系 | 2013-2017')

        self.save_document(doc, 'edge-13-multiple-positions.docx')

    def generate_edge_14_freelancer(self):
        """自由职业者简历"""
        doc = self.create_document()

        title = doc.add_heading('Alex Zhang', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Freelance Developer | +86-139-0013-9000 | alex@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'PROJECT EXPERIENCE', 2)
        doc.add_paragraph('E-commerce Platform | Client: Startup A | 2023.06-2023.09')
        doc.add_paragraph('• Built full-stack web application')
        doc.add_paragraph()

        doc.add_paragraph('Mobile App | Client: Company B | 2023.03-2023.05')
        doc.add_paragraph('• Developed iOS app')
        doc.add_paragraph()

        doc.add_paragraph('Website Redesign | Client: Agency C | 2023.01-2023.02')
        doc.add_paragraph('• Frontend development')
        doc.add_paragraph()

        doc.add_paragraph('API Integration | Client: Enterprise D | 2022.10-2022.12')
        doc.add_paragraph('• Backend services')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION', 2)
        doc.add_paragraph('Shanghai University | Computer Science | 2015-2019')

        self.save_document(doc, 'edge-14-freelancer.docx')

    def generate_edge_15_long_resume(self):
        """超长简历（5页+）"""
        doc = self.create_document()

        title = doc.add_heading('资深专家', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('电话: 13700137777 | 邮箱: expert@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        for i in range(12):
            doc.add_paragraph(f'公司{i+1} | 职位{i+1} | {2022-i}.01-{2023-i}.12')
            doc.add_paragraph(f'• 工作内容描述 {i+1}')
            doc.add_paragraph(f'• 项目经验 {i+1}')
            doc.add_paragraph(f'• 技术栈 {i+1}')
            doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('某大学 | 博士 | 2008-2012')
        doc.add_paragraph('某大学 | 硕士 | 2006-2008')
        doc.add_paragraph('某大学 | 本科 | 2002-2006')

        self.save_document(doc, 'edge-15-long-resume.docx')

    def generate_edge_16_short_resume(self):
        """超短简历（半页）"""
        doc = self.create_document()

        title = doc.add_heading('简短', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('电话: 13600136666 | 邮箱: short@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('某公司 | 职位 | 2022-2023')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('某大学 | 本科 | 2018-2022')

        self.save_document(doc, 'edge-16-short-resume.docx')

    def generate_edge_17_special_chars(self):
        """特殊字符简历"""
        doc = self.create_document()

        title = doc.add_heading('特殊@符号', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('电话: 13500135555 | 邮箱: special@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('A&B科技(集团)有限公司 | 高级工程师(后端) | 2020-2023')
        doc.add_paragraph('• 负责C++/Python开发')
        doc.add_paragraph()

        doc.add_paragraph('X-Y互联网 | 全栈工程师 | 2018-2020')
        doc.add_paragraph('• 使用Node.js/React')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('某大学(985) | 计算机 | 2014-2018')

        self.save_document(doc, 'edge-17-special-chars.docx')

    def generate_edge_18_unstructured(self):
        """无结构文本"""
        doc = self.create_document()

        title = doc.add_heading('无结构简历', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        doc.add_paragraph('我叫张三，电话13400134444，邮箱zhangsan@example.com。我在2020年到2023年在某科技公司工作，担任软件工程师。主要负责后端开发工作，使用Java和Python。之前在2018年到2020年在另一家公司做过类似的工作。我2018年毕业于某大学计算机专业，获得本科学位。')

        self.save_document(doc, 'edge-18-unstructured.docx')

    def generate_all_edge_cases(self):
        """生成所有边缘情况简历"""
        print("开始生成第二阶段边缘情况测试集...")
        self.generate_edge_11_missing_info()
        self.generate_edge_12_date_formats()
        self.generate_edge_13_multiple_positions()
        self.generate_edge_14_freelancer()
        self.generate_edge_15_long_resume()
        self.generate_edge_16_short_resume()
        self.generate_edge_17_special_chars()
        self.generate_edge_18_unstructured()
        print(f"\n✓ 第二阶段完成！共生成 8 份边缘情况简历")


if __name__ == '__main__':
    generator = EdgeCaseGenerator('test-resumes/phase2-edge-cases')
    generator.generate_all_edge_cases()

