#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成测试简历数据集
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class ResumeGenerator:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(self):
        """创建新文档"""
        doc = Document()
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        return doc

    def add_heading(self, doc, text, level=1):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading

    def add_paragraph(self, doc, text, bold=False):
        """添加段落"""
        p = doc.add_paragraph(text)
        if bold:
            p.runs[0].bold = True
        return p

    def save_document(self, doc, filename):
        """保存文档"""
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        print(f"✓ 已生成: {filename}")

    def generate_cn_01_traditional_table(self):
        """中文传统表格式简历"""
        doc = self.create_document()

        # 标题
        title = doc.add_heading('个人简历', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息表格
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        cells = [
            ['姓名', '张伟', '性别', '男'],
            ['出生日期', '1990-05-15', '民族', '汉族'],
            ['电话', '13800138000', '邮箱', 'zhangwei@example.com'],
            ['微信', 'zhangwei_wx', '学历', '本科'],
            ['现居地', '北京市朝阳区', '婚姻状况', '已婚'],
            ['求职意向', '软件工程师', '期望薪资', '15-20K']
        ]

        for i, row_data in enumerate(cells):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = cell_text

        doc.add_paragraph()

        # 工作经历
        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('北京科技有限公司 | 高级软件工程师 | 2020.01-2022.12')
        doc.add_paragraph('• 负责公司核心产品的后端开发，使用Java Spring Boot框架')
        doc.add_paragraph('• 优化数据库查询性能，将响应时间降低40%')
        doc.add_paragraph('• 带领3人团队完成微服务架构改造')
        doc.add_paragraph()

        doc.add_paragraph('上海互联网公司 | 软件工程师 | 2018.03-2019.12')
        doc.add_paragraph('• 参与电商平台开发，负责订单模块')
        doc.add_paragraph('• 实现分布式缓存方案，提升系统并发能力')
        doc.add_paragraph()

        doc.add_paragraph('深圳创业公司 | 初级工程师 | 2016.07-2018.02')
        doc.add_paragraph('• 负责移动端API开发')
        doc.add_paragraph('• 参与需求分析和技术方案设计')
        doc.add_paragraph()

        # 教育背景
        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('清华大学 | 计算机科学与技术 | 本科 | 2012.09-2016.06')
        doc.add_paragraph('北京理工大学 | 软件工程 | 硕士 | 2016.09-2018.06')

        self.save_document(doc, 'cn-01-traditional-table.docx')

    def generate_cn_02_modern_tech(self):
        """中文现代简洁式简历（互联网行业）"""
        doc = self.create_document()

        # 标题和基本信息
        title = doc.add_heading('李明', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph('Python后端工程师')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact = doc.add_paragraph('电话: 13900139000 | 邮箱: liming@example.com | GitHub: github.com/liming')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # 技术栈
        self.add_heading(doc, '技术栈', 2)
        doc.add_paragraph('• 编程语言: Python, Go, JavaScript')
        doc.add_paragraph('• 框架: Django, Flask, FastAPI, Vue.js')
        doc.add_paragraph('• 数据库: MySQL, PostgreSQL, Redis, MongoDB')
        doc.add_paragraph('• 工具: Docker, Kubernetes, Git, Jenkins')
        doc.add_paragraph()

        # 工作经历
        self.add_heading(doc, '工作经历', 2)

        doc.add_paragraph('字节跳动 | 高级后端工程师 | 2021年3月-2024年2月')
        doc.add_paragraph('• 负责推荐系统后端服务开发，日均处理10亿+请求')
        doc.add_paragraph('• 设计并实现分布式任务调度系统，支持百万级任务并发')
        doc.add_paragraph('• 优化服务性能，P99延迟降低60%')
        doc.add_paragraph()

        doc.add_paragraph('美团 | 后端工程师 | 2019年7月-2021年2月')
        doc.add_paragraph('• 参与外卖配送系统开发，负责订单分配模块')
        doc.add_paragraph('• 实现实时数据同步方案，确保数据一致性')
        doc.add_paragraph()

        doc.add_paragraph('腾讯 | 后端开发实习生 | 2018年12月-2019年6月')
        doc.add_paragraph('• 参与微信支付后台开发')
        doc.add_paragraph('• 编写单元测试，代码覆盖率达到85%')
        doc.add_paragraph()

        doc.add_paragraph('阿里巴巴 | 后端开发实习生 | 2018年6月-2018年11月')
        doc.add_paragraph('• 参与淘宝商品搜索优化项目')
        doc.add_paragraph()

        # 教育背景
        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('北京大学 | 计算机科学与技术 | 本科 | 2015.09-2019.06')

        self.save_document(doc, 'cn-02-modern-tech.docx')

    def generate_cn_03_two_column(self):
        """中文双栏式简历"""
        doc = self.create_document()

        title = doc.add_heading('王芳', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('产品经理 | 13700137000 | wangfang@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # 技能与证书
        self.add_heading(doc, '专业技能', 2)
        doc.add_paragraph('• Axure RP ★★★★★')
        doc.add_paragraph('• Figma ★★★★☆')
        doc.add_paragraph('• SQL ★★★★☆')
        doc.add_paragraph('• Python ★★★☆☆')
        doc.add_paragraph()

        self.add_heading(doc, '证书', 2)
        doc.add_paragraph('• PMP项目管理专业人士认证')
        doc.add_paragraph('• 产品经理国际资格认证(NPDP)')
        doc.add_paragraph()

        # 工作经历
        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('京东 | 高级产品经理 | 2020.06-至今')
        doc.add_paragraph('• 负责电商平台用户增长产品设计')
        doc.add_paragraph('• 主导会员体系改版，用户留存率提升25%')
        doc.add_paragraph()

        doc.add_paragraph('小米 | 产品经理 | 2018.08-2020.05')
        doc.add_paragraph('• 负责智能家居APP产品规划')
        doc.add_paragraph()

        doc.add_paragraph('百度 | 产品助理 | 2016.07-2018.07')
        doc.add_paragraph('• 协助搜索产品优化')
        doc.add_paragraph()

        # 教育背景
        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('复旦大学 | 工商管理 | 本科 | 2012.09-2016.06')
        doc.add_paragraph('上海交通大学 | MBA | 硕士 | 2018.09-2020.06')

        self.save_document(doc, 'cn-03-two-column.docx')

    def generate_cn_04_finance(self):
        """中文金融行业简历"""
        doc = self.create_document()

        title = doc.add_heading('陈强', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('金融分析师 | 13600136000 | chenqiang@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '工作经历', 2)
        doc.add_paragraph('中信证券 | 高级分析师 | 2021.01-至今')
        doc.add_paragraph('• 负责A股市场行业研究，覆盖科技、医药板块')
        doc.add_paragraph('• 完成IPO项目3个，融资总额超50亿元')
        doc.add_paragraph()

        doc.add_paragraph('招商银行 | 投资顾问 | 2019.03-2020.12')
        doc.add_paragraph('• 管理高净值客户资产组合，规模达2亿元')
        doc.add_paragraph()

        doc.add_paragraph('平安保险 | 理财经理 | 2017.07-2019.02')
        doc.add_paragraph('• 年均销售业绩1500万元')
        doc.add_paragraph()

        doc.add_paragraph('工商银行 | 客户经理 | 2015.08-2017.06')
        doc.add_paragraph('• 维护企业客户关系')
        doc.add_paragraph()

        doc.add_paragraph('建设银行 | 实习生 | 2015.03-2015.07')
        doc.add_paragraph('• 协助客户开户业务')
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('中央财经大学 | 金融学 | 本科 | 2011.09-2015.06')
        doc.add_paragraph('清华大学 | 金融硕士 | 硕士 | 2015.09-2017.06')

        self.save_document(doc, 'cn-04-finance.docx')

    def generate_cn_05_fresh_graduate(self):
        """中文应届生简历"""
        doc = self.create_document()

        title = doc.add_heading('刘洋', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('应届毕业生 | 13500135000 | liuyang@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, '教育背景', 2)
        doc.add_paragraph('浙江大学 | 软件工程 | 本科 | 2020.09-2024.06（在校）')
        doc.add_paragraph('GPA: 3.8/4.0 | 专业排名: 5/120')
        doc.add_paragraph()

        self.add_heading(doc, '实习经历', 2)
        doc.add_paragraph('网易 | 前端开发实习生 | 2023.06-2023.12')
        doc.add_paragraph('• 参与游戏官网开发，使用React框架')
        doc.add_paragraph('• 优化页面加载速度，首屏时间减少30%')
        doc.add_paragraph()

        doc.add_paragraph('华为 | 软件测试实习生 | 2022.12-2023.05')
        doc.add_paragraph('• 编写自动化测试脚本')
        doc.add_paragraph()

        self.add_heading(doc, '项目经验', 2)
        doc.add_paragraph('在线教育平台（毕业设计）')
        doc.add_paragraph('• 使用Vue.js + Spring Boot开发')
        doc.add_paragraph()

        self.save_document(doc, 'cn-05-fresh-graduate.docx')

    def generate_en_06_us_tech(self):
        """US Format - Tech Resume"""
        doc = self.create_document()

        title = doc.add_heading('John Smith', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Software Engineer | +1-555-0123 | john.smith@example.com | linkedin.com/in/johnsmith').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'SKILLS', 2)
        doc.add_paragraph('Languages: Java, Python, JavaScript, TypeScript, SQL')
        doc.add_paragraph('Frameworks: Spring Boot, React, Node.js, Django')
        doc.add_paragraph('Tools: AWS, Docker, Kubernetes, Git, Jenkins')
        doc.add_paragraph()

        self.add_heading(doc, 'EXPERIENCE', 2)
        doc.add_paragraph('Google | Senior Software Engineer | Jan 2020 - Dec 2022')
        doc.add_paragraph('• Developed scalable microservices handling 1M+ requests per day')
        doc.add_paragraph('• Reduced API latency by 45% through optimization')
        doc.add_paragraph()

        doc.add_paragraph('Amazon | Software Engineer | Mar 2018 - Dec 2019')
        doc.add_paragraph('• Built e-commerce features for Prime platform')
        doc.add_paragraph()

        doc.add_paragraph('Microsoft | Software Engineer | Jul 2016 - Feb 2018')
        doc.add_paragraph('• Worked on Azure cloud services')
        doc.add_paragraph()

        doc.add_paragraph('Facebook | Software Engineering Intern | Jun 2015 - Aug 2015')
        doc.add_paragraph('• Developed internal tools')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION', 2)
        doc.add_paragraph('Stanford University | B.S. Computer Science | Sep 2012 - Jun 2016')

        self.save_document(doc, 'en-06-us-tech.docx')

    def generate_en_07_uk_cv(self):
        """UK CV Format"""
        doc = self.create_document()

        title = doc.add_heading('Emily Johnson', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Marketing Manager | +44-20-1234-5678 | emily.johnson@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'PERSONAL STATEMENT', 2)
        doc.add_paragraph('Experienced marketing professional with 8+ years in digital marketing and brand management.')
        doc.add_paragraph()

        self.add_heading(doc, 'EMPLOYMENT HISTORY', 2)
        doc.add_paragraph('Unilever | Senior Marketing Manager | January 2020 - December 2023')
        doc.add_paragraph('• Led brand strategy for 5 product lines')
        doc.add_paragraph()

        doc.add_paragraph('Procter & Gamble | Marketing Manager | March 2018 - December 2019')
        doc.add_paragraph('• Managed £2M marketing budget')
        doc.add_paragraph()

        doc.add_paragraph('L\'Oreal | Brand Executive | July 2016 - February 2018')
        doc.add_paragraph('• Coordinated product launches')
        doc.add_paragraph()

        doc.add_paragraph('Nestle | Marketing Assistant | September 2014 - June 2016')
        doc.add_paragraph('• Supported campaign execution')
        doc.add_paragraph()

        doc.add_paragraph('Coca-Cola | Marketing Intern | June 2014 - August 2014')
        doc.add_paragraph('• Assisted with market research')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION', 2)
        doc.add_paragraph('University of Oxford | BA Marketing | October 2010 - June 2014')
        doc.add_paragraph('London Business School | MBA | September 2016 - June 2018')

        self.save_document(doc, 'en-07-uk-cv.docx')

    def generate_en_08_creative(self):
        """Creative Industry Resume"""
        doc = self.create_document()

        title = doc.add_heading('Sarah Chen', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('UX/UI Designer | +1-555-0456 | sarah.chen@example.com | portfolio.sarahchen.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'EXPERIENCE', 2)
        doc.add_paragraph('Adobe | Senior UX Designer | Mar 2021 - Present')
        doc.add_paragraph('• Design user experiences for Creative Cloud apps')
        doc.add_paragraph()

        doc.add_paragraph('Airbnb | UX Designer | Jun 2019 - Feb 2021')
        doc.add_paragraph('• Redesigned booking flow, increased conversion by 18%')
        doc.add_paragraph()

        doc.add_paragraph('Spotify | UI Designer | Aug 2017 - May 2019')
        doc.add_paragraph('• Created visual designs for mobile app')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION', 2)
        doc.add_paragraph('Rhode Island School of Design | BFA Graphic Design | Sep 2013 - May 2017')

        self.save_document(doc, 'en-08-creative.docx')

    def generate_en_09_executive(self):
        """Executive Resume"""
        doc = self.create_document()

        title = doc.add_heading('Michael Anderson', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Chief Technology Officer | +1-555-0789 | m.anderson@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'EXECUTIVE SUMMARY', 2)
        doc.add_paragraph('Technology executive with 15+ years leading engineering teams and driving digital transformation.')
        doc.add_paragraph()

        self.add_heading(doc, 'PROFESSIONAL EXPERIENCE', 2)
        doc.add_paragraph('TechCorp Inc. | Chief Technology Officer | Jan 2020 - Present')
        doc.add_paragraph('• Lead 200+ engineering team, $50M budget')
        doc.add_paragraph()

        doc.add_paragraph('DataSystems Ltd. | VP Engineering | Mar 2017 - Dec 2019')
        doc.add_paragraph('• Built engineering organization from 20 to 100 people')
        doc.add_paragraph()

        doc.add_paragraph('CloudTech | Director of Engineering | Jul 2014 - Feb 2017')
        doc.add_paragraph('• Managed 5 engineering teams')
        doc.add_paragraph()

        doc.add_paragraph('SoftwareCo | Engineering Manager | Jan 2012 - Jun 2014')
        doc.add_paragraph('• Led backend development team')
        doc.add_paragraph()

        doc.add_paragraph('StartupXYZ | Senior Engineer | Mar 2010 - Dec 2011')
        doc.add_paragraph('• Early engineer at Series A startup')
        doc.add_paragraph()

        doc.add_paragraph('IBM | Software Engineer | Jul 2008 - Feb 2010')
        doc.add_paragraph('• Enterprise software development')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION', 2)
        doc.add_paragraph('MIT | B.S. Computer Science | Sep 2004 - Jun 2008')
        doc.add_paragraph('Harvard Business School | MBA | Sep 2012 - Jun 2014')

        self.save_document(doc, 'en-09-executive.docx')

    def generate_en_10_bilingual(self):
        """Bilingual Resume (中英混合)"""
        doc = self.create_document()

        title = doc.add_heading('王磊 (David Wang)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('International Business Manager | +86-138-0013-8888 | david.wang@example.com').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        self.add_heading(doc, 'WORK EXPERIENCE / 工作经历', 2)
        doc.add_paragraph('Huawei Technologies | Senior Product Manager | Jan 2020 - Present')
        doc.add_paragraph('华为技术有限公司 | 高级产品经理 | 2020年1月-至今')
        doc.add_paragraph('• Led international market expansion in Europe')
        doc.add_paragraph('• 负责欧洲市场拓展，年销售额达5000万美元')
        doc.add_paragraph()

        doc.add_paragraph('Alibaba Group | Product Manager | Mar 2018 - Dec 2019')
        doc.add_paragraph('阿里巴巴集团 | 产品经理 | 2018年3月-2019年12月')
        doc.add_paragraph('• Managed cross-border e-commerce platform')
        doc.add_paragraph()

        doc.add_paragraph('Tencent | Business Analyst | Jul 2016 - Feb 2018')
        doc.add_paragraph('腾讯 | 商业分析师 | 2016年7月-2018年2月')
        doc.add_paragraph('• Analyzed market trends for international expansion')
        doc.add_paragraph()

        self.add_heading(doc, 'EDUCATION / 教育背景', 2)
        doc.add_paragraph('Peking University | Bachelor of Business Administration | Sep 2012 - Jun 2016')
        doc.add_paragraph('北京大学 | 工商管理学士 | 2012年9月-2016年6月')
        doc.add_paragraph()
        doc.add_paragraph('University of Cambridge | Master of Business Administration | Sep 2016 - Jun 2017')
        doc.add_paragraph('剑桥大学 | 工商管理硕士 | 2016年9月-2017年6月')

        self.save_document(doc, 'en-10-bilingual.docx')

    def generate_all_phase1(self):
        """生成第一阶段所有简历"""
        print("开始生成第一阶段核心测试集...")
        self.generate_cn_01_traditional_table()
        self.generate_cn_02_modern_tech()
        self.generate_cn_03_two_column()
        self.generate_cn_04_finance()
        self.generate_cn_05_fresh_graduate()
        self.generate_en_06_us_tech()
        self.generate_en_07_uk_cv()
        self.generate_en_08_creative()
        self.generate_en_09_executive()
        self.generate_en_10_bilingual()
        print(f"\n✓ 第一阶段完成！共生成 10 份简历")


if __name__ == '__main__':
    generator = ResumeGenerator('test-resumes/phase1-core')
    generator.generate_all_phase1()

